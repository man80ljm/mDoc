import os
import tkinter as tk
from tkinter import filedialog, ttk, messagebox
from docx import Document
from docx.shared import Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import sys
import re
import time
from datetime import datetime

class FolderToWordApp:
    def __init__(self, root):
        self.root = root
        self.root.title("佐证材料自动生成工具 - 强化版")
        self.root.geometry("600x400")
        self.root.resizable(False, False)

        self.folder_path = tk.StringVar()
        self.debug_info = []

        self.create_widgets()

    def create_widgets(self):
        # 标题
        title_label = tk.Label(self.root, text="请先选择需要处理的文件夹", font=("Microsoft YaHei", 12))
        title_label.pack(pady=10)

        # 按钮框架
        btn_frame = tk.Frame(self.root)
        btn_frame.pack(pady=10)

        self.load_btn = tk.Button(btn_frame, text="1. 加载文件夹", width=18, height=2,
                                  bg="#f0f0f0", font=("Microsoft YaHei", 10))
        self.load_btn.grid(row=0, column=0, padx=15)

        self.generate_btn = tk.Button(btn_frame, text="2. 生成佐证材料.docx", width=25, height=2,
                                      bg="#28a745", fg="white", font=("Microsoft YaHei", 10, "bold"),
                                      state=tk.DISABLED)
        self.generate_btn.grid(row=0, column=1, padx=15)

        self.load_btn.config(command=self.load_folder)
        self.generate_btn.config(command=self.generate_word)

        # 页边距设置
        path_frame = tk.Frame(self.root)
        path_frame.pack(pady=10, fill=tk.X, padx=40)
        tk.Label(path_frame, text="页面布局（页边距）", font=("Microsoft YaHei", 10)).pack(anchor=tk.W)
        
        self.margin_var = tk.StringVar(value="窄")
        margin_combo = ttk.Combobox(path_frame, textvariable=self.margin_var, state="readonly",
                                    font=("Microsoft YaHei", 10), width=30)
        margin_combo['values'] = ("窄", "普通", "适中", "宽")
        margin_combo.pack(pady=5)

        # 当前选中路径显示
        self.path_label = tk.Label(self.root, text="未选择文件夹", fg="gray", font=("Microsoft YaHei", 9),
                                   wraplength=520, justify="left")
        self.path_label.pack(pady=10)

        # 实时调试信息显示区域
        debug_frame = tk.Frame(self.root)
        debug_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        tk.Label(debug_frame, text="实时调试信息：", font=("Microsoft YaHei", 10)).pack(anchor=tk.W)
        
        # 创建文本框和滚动条
        text_frame = tk.Frame(debug_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        self.debug_text = tk.Text(text_frame, height=8, font=("Consolas", 9))
        scrollbar = tk.Scrollbar(text_frame, orient="vertical", command=self.debug_text.yview)
        self.debug_text.configure(yscrollcommand=scrollbar.set)
        
        self.debug_text.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # 进度条
        self.progress = ttk.Progressbar(self.root, mode='indeterminate')
        self.progress.pack(fill=tk.X, padx=40, pady=10)

    def add_debug_info(self, message):
        """添加调试信息并实时显示"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        full_message = f"[{timestamp}] {message}"
        self.debug_info.append(full_message)
        
        # 实时更新文本框
        self.debug_text.insert(tk.END, full_message + "\n")
        self.debug_text.see(tk.END)
        self.root.update()
        print(full_message)

    def load_folder(self):
        folder = filedialog.askdirectory(title="请选择包含佐证材料的根文件夹")
        if folder:
            self.folder_path.set(folder)
            self.path_label.config(text=folder, fg="black")
            self.generate_btn.config(state=tk.NORMAL)
            
            # 清空调试信息
            self.debug_text.delete(1.0, tk.END)
            self.debug_info.clear()
            self.add_debug_info(f"已选择文件夹: {folder}")

    def get_margin_cm(self):
        margins = {
            "窄": 1.27,
            "普通": 2.54,
            "适中": 1.9,
            "宽": 3.17
        }
        return margins.get(self.margin_var.get(), 1.27)

    def is_image_file(self, filename):
        """检查文件是否为图片"""
        image_extensions = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tif', '.tiff', '.webp'}
        return os.path.splitext(filename.lower())[1] in image_extensions

    def scan_folder_structure(self, folder_path, level=0):
        """预扫描文件夹结构，找出所有叶子节点"""
        folder_name = os.path.basename(folder_path)
        indent = "  " * level
        
        try:
            items = os.listdir(folder_path)
        except Exception as e:
            self.add_debug_info(f"{indent}无法读取: {folder_name} - {e}")
            return
        
        subfolders = []
        images = []
        
        for item in items:
            item_path = os.path.join(folder_path, item)
            if os.path.isdir(item_path):
                subfolders.append(item_path)
            elif self.is_image_file(item):
                images.append(item_path)
        
        # 判断是否为叶子节点
        is_leaf = len(subfolders) == 0
        self.add_debug_info(f"{indent}{folder_name}: {len(images)}图片, {len(subfolders)}子文件夹, {'叶子节点' if is_leaf else '非叶子节点'}")
        
        # 递归扫描子文件夹
        for subfolder in sorted(subfolders):
            self.scan_folder_structure(subfolder, level + 1)

    def generate_word(self):
        if not self.folder_path.get():
            messagebox.showwarning("警告", "请先加载文件夹！")
            return

        self.add_debug_info("=" * 50)
        self.add_debug_info("开始预扫描文件夹结构...")
        
        # 先扫描整个文件夹结构
        self.scan_folder_structure(self.folder_path.get())
        
        self.add_debug_info("=" * 50)
        self.add_debug_info("开始生成Word文档...")
        self.progress.start(10)
        self.generate_btn.config(state=tk.DISABLED)

        try:
            doc = Document()
            self.setup_page_margins(doc)
            self.add_debug_info("已创建新Word文档并设置页边距")

            root_folder = self.folder_path.get()
            root_name = os.path.basename(root_folder)

            # 添加总标题
            title = doc.add_heading(root_name, level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self.add_debug_info(f"已添加总标题: {root_name}")

            # 递归生成目录和图片
            self.process_folder(doc, root_folder, level=1)

            # 保存文档
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            save_path = os.path.join(root_folder, f"佐证材料_{timestamp}.docx")
            doc.save(save_path)
            self.add_debug_info(f"文档已保存: {save_path}")

            messagebox.showinfo("成功", f"佐证材料生成成功！\n保存路径：\n{save_path}")

        except Exception as e:
            error_msg = f"生成失败: {str(e)}"
            self.add_debug_info(error_msg)
            messagebox.showerror("错误", error_msg)
        finally:
            self.progress.stop()
            self.generate_btn.config(state=tk.NORMAL)

    def setup_page_margins(self, doc):
        margin_cm = self.get_margin_cm()
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(margin_cm)
            section.bottom_margin = Cm(margin_cm)
            section.left_margin = Cm(margin_cm)
            section.right_margin = Cm(margin_cm)

    def has_subfolders(self, folder_path):
        """检查文件夹是否有子文件夹"""
        try:
            items = os.listdir(folder_path)
            for item in items:
                if os.path.isdir(os.path.join(folder_path, item)):
                    return True
            return False
        except:
            return False

    def process_folder(self, doc, folder_path, level=1):
        folder_name = os.path.basename(folder_path)
        self.add_debug_info(f"{'  ' * level}处理文件夹: {folder_name} (层级{level})")
        
        try:
            items = os.listdir(folder_path)
        except Exception as e:
            self.add_debug_info(f"{'  ' * level}无法读取文件夹: {e}")
            return
            
        subfolders = []
        images = []

        # 分类文件和文件夹
        for item in items:
            item_path = os.path.join(folder_path, item)
            try:
                if os.path.isdir(item_path):
                    subfolders.append(item_path)
                elif os.path.isfile(item_path) and self.is_image_file(item):
                    images.append(item_path)
            except Exception as e:
                self.add_debug_info(f"{'  ' * level}处理项目出错 {item}: {e}")
                continue

        # 自然排序
        def natural_key(p): 
            return [int(c) if c.isdigit() else c.lower() for c in re.split(r'(\d+)', os.path.basename(p))]
        
        subfolders = sorted(subfolders, key=natural_key)
        images = sorted(images, key=natural_key)

        # 添加标题
        if level <= 9:
            doc.add_heading(folder_name, level=level)
        else:
            p = doc.add_paragraph(folder_name)
            p.style = 'Heading 9'

        # 关键改进：明确的叶子节点判断
        is_leaf_node = len(subfolders) == 0
        
        self.add_debug_info(f"{'  ' * level}{folder_name}: {len(images)}图片, {len(subfolders)}子文件夹 -> {'叶子节点' if is_leaf_node else '非叶子节点'}")

        # 只有叶子节点才插入图片
        if is_leaf_node and images:
            self.add_debug_info(f"{'  ' * level}→ 开始插入图片到叶子节点")
            inserted_count = self.insert_images_enhanced(doc, images, level)
            self.add_debug_info(f"{'  ' * level}→ 插入完成: {inserted_count}/{len(images)}")
        elif not is_leaf_node:
            self.add_debug_info(f"{'  ' * level}→ 非叶子节点，跳过图片插入")
        elif is_leaf_node and not images:
            self.add_debug_info(f"{'  ' * level}→ 叶子节点但无图片")
        
        # 递归处理子文件夹
        for sub in subfolders:
            self.process_folder(doc, sub, level + 1)

    def insert_images_enhanced(self, doc, image_paths, level):
        """终极修复版：预先调整图片大小，彻底解决插入失败问题"""
        margin_cm = self.get_margin_cm()
        usable_width_cm = 21.0 - 2 * margin_cm
        max_page_height_cm = 29.7 - 2 * margin_cm - 1
        
        # 图片缩放比例：90%
        scale_factor = 0.9

        inserted_count = 0
        indent = "  " * (level + 1)

        for i, img_path in enumerate(image_paths, 1):
            filename = os.path.basename(img_path)
            self.add_debug_info(f"{indent}[{i}/{len(image_paths)}] 处理: {filename}")

            try:
                if not os.path.exists(img_path) or os.path.getsize(img_path) == 0:
                    self.add_debug_info(f"{indent}  ✗ 文件不存在或为空")
                    continue

                # 用PIL打开图片
                with Image.open(img_path) as img:
                    width_px, height_px = img.size
                    mode = img.mode
                    self.add_debug_info(f"{indent}  尺寸: {width_px}×{height_px}  模式: {mode}")

                    # 转换为RGB（如果需要）
                    if mode in ('RGBA', 'P', 'LA', 'PA'):
                        img = img.convert('RGB')
                        self.add_debug_info(f"{indent}  已转换为RGB")
                    
                    # 关键改进：预先调整图片大小
                    # 目标：让最长边不超过2000px
                    max_dimension = 2000
                    if max(width_px, height_px) > max_dimension:
                        ratio = max_dimension / max(width_px, height_px)
                        new_width = int(width_px * ratio)
                        new_height = int(height_px * ratio)
                        img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        self.add_debug_info(f"{indent}  已调整尺寸: {new_width}×{new_height}")
                        width_px, height_px = new_width, new_height
                    
                    # 保存为临时文件
                    temp_path = img_path + '_temp_resized.jpg'
                    img.save(temp_path, 'JPEG', quality=95)
                    img_path_to_use = temp_path

                # 创建居中段落
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run()

                # 简化策略：直接根据方向选择尺寸参数
                success = False
                try:
                    if height_px > width_px:
                        # 竖版图限高
                        run.add_picture(img_path_to_use, height=Cm(max_page_height_cm * scale_factor))
                        self.add_debug_info(f"{indent}  ✓ 成功：竖版图限高90%")
                    else:
                        # 横版图限宽
                        run.add_picture(img_path_to_use, width=Cm(usable_width_cm * 0.88 * scale_factor))
                        self.add_debug_info(f"{indent}  ✓ 成功：横版图限宽79%")
                    success = True
                    inserted_count += 1
                except Exception as e:
                    self.add_debug_info(f"{indent}  ✗ 插入失败: {str(e)[:80]}")
                    # 删除失败的空段落
                    doc._element.body.remove(p._element)

                # 清理临时文件
                if os.path.exists(temp_path):
                    try:
                        os.remove(temp_path)
                    except:
                        pass

            except Exception as e:
                self.add_debug_info(f"{indent}  ✗ 处理异常: {str(e)[:80]}")
                # 临时文件清理
                if 'temp_path' in locals() and os.path.exists(temp_path):
                    try:
                        os.remove(temp_path)
                    except:
                        pass
                continue

        return inserted_count

if __name__ == "__main__":
    root = tk.Tk()
    app = FolderToWordApp(root)
    root.mainloop()